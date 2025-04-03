import io
import docx
import pypdf
import re
import os
from typing import Dict, Any, Optional

class FileExtractor:
    """Klasse zur Extraktion von Text aus verschiedenen Dateiformaten (PDF, DOCX, TXT)"""
    
    def __init__(self):
        self.supported_formats = {
            ".pdf": self.extract_from_pdf,
            ".docx": self.extract_from_docx,
            ".doc": self.extract_from_doc,
            ".txt": self.extract_from_txt
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extrahiert Text aus einer Datei basierend auf ihrem Format"""
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            return {
                "success": False,
                "error": f"Nicht unterstütztes Dateiformat: {file_ext}",
                "text": ""
            }
            
        try:
            extractor = self.supported_formats[file_ext]
            text = extractor(file_path)
            
            return {
                "success": True,
                "text": text,
                "format": file_ext[1:],  # Entferne den Punkt vom Format
                "word_count": len(text.split())
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            }
    
    def extract_from_pdf(self, file_path: str) -> str:
        """Extrahiert Text aus einer PDF-Datei"""
        
        text = ""
        
        with open(file_path, "rb") as file:
            pdf = pypdf.PdfReader(file)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        # Bereinige Text
        text = self.clean_text(text)
        return text
    
    def extract_from_docx(self, file_path: str) -> str:
        """Extrahiert Text aus einer DOCX-Datei"""
        
        doc = docx.Document(file_path)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        # Bereinige Text
        text = self.clean_text(text)
        return text
    
    def extract_from_doc(self, file_path: str) -> str:
        """Extrahiert Text aus einer älteren DOC-Datei (benötigt externe Konvertierung)"""
        
        # Hinweis: Dies ist ein Platzhalter. Die tatsächliche DOC-Extraktion könnte
        # externe Tools wie Antiword, textract oder eine Konvertierung zu DOCX erfordern
        
        return "Die Extraktion aus DOC-Dateien erfordert zusätzliche Libraries.\n" + \
               "Bitte konvertieren Sie die Datei in DOCX oder PDF."
    
    def extract_from_txt(self, file_path: str) -> str:
        """Extrahiert Text aus einer Textdatei"""
        
        with open(file_path, "r", encoding="utf-8", errors="replace") as file:
            text = file.read()
            
        # Bereinige Text
        text = self.clean_text(text)
        return text
    
    def clean_text(self, text: str) -> str:
        """Bereinigt extrahierten Text von Formatierungsproblemen"""
        
        # Ersetze multiple Leerzeichen
        text = re.sub(r'\s+', ' ', text)
        
        # Entferne Seitenumbrüche und andere Sonderzeichen
        text = re.sub(r'\f', '\n\n', text)
        
        # Entferne ungewöhnliche Unicode-Zeichen
        text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]', '', text)
        
        # Normalisiere Zeilenumbrüche
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def extract_from_stream(self, file_stream: io.BytesIO, file_name: str) -> Dict[str, Any]:
        """Extrahiert Text aus einem Dateistrom anstatt einer Datei auf der Festplatte"""
        
        file_ext = os.path.splitext(file_name)[1].lower()
        
        if file_ext not in self.supported_formats:
            return {
                "success": False,
                "error": f"Nicht unterstütztes Dateiformat: {file_ext}",
                "text": ""
            }
            
        try:
            if file_ext == ".pdf":
                pdf = pypdf.PdfReader(file_stream)
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
            elif file_ext == ".docx":
                doc = docx.Document(file_stream)
                text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                
            elif file_ext == ".txt":
                text = file_stream.read().decode("utf-8", errors="replace")
                
            else:
                # Sollte nicht hierher gelangen, da wir bereits das Format geprüft haben
                return {
                    "success": False,
                    "error": f"Unerwarteter Fehler beim Verarbeiten des Formats {file_ext}",
                    "text": ""
                }
            
            # Text bereinigen
            text = self.clean_text(text)
            
            return {
                "success": True,
                "text": text,
                "format": file_ext[1:],
                "word_count": len(text.split())
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": ""
            } 