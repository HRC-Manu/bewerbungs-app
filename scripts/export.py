#!/usr/bin/env python3
"""
Document export utility for the Job Application Assistant.
Converts generated documents to Word and PDF formats.
"""

import os
import sys
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import argparse
from docx2pdf import convert

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentExporter:
    """Handles the export of documents to various formats."""
    
    def __init__(self, input_dir: str = 'input', output_dir: str = 'output'):
        """Initialize the exporter with input and output directories."""
        self.input_dir = input_dir
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def html_to_text(self, html_content: str) -> str:
        """Convert HTML content to plain text while preserving structure."""
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text()
    
    def create_word_document(self) -> Document:
        """Create and format a new Word document."""
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)
        
        return doc
        
    def export_to_word(self) -> Dict[str, Any]:
        """Export the documents to Word format."""
        try:
            # Create a new Word document
            doc = self.create_word_document()
            
            # Add cover letter
            with open(os.path.join(self.input_dir, 'cover_letter.html'), 'r', encoding='utf-8') as f:
                cover_letter_html = f.read()
            
            cover_letter_text = self.html_to_text(cover_letter_html)
            
            # Format cover letter
            heading = doc.add_heading('Anschreiben', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add cover letter text with proper formatting
            paragraphs = cover_letter_text.split('\n\n')
            for p_text in paragraphs:
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.paragraph_format.space_after = Pt(12)
            
            # Add page break
            doc.add_page_break()
            
            # Add resume
            with open(os.path.join(self.input_dir, 'resume.html'), 'r', encoding='utf-8') as f:
                resume_html = f.read()
            
            resume_text = self.html_to_text(resume_html)
            
            # Format resume
            heading = doc.add_heading('Lebenslauf', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add resume text with proper formatting
            paragraphs = resume_text.split('\n\n')
            for p_text in paragraphs:
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.paragraph_format.space_after = Pt(12)
            
            # Save the document
            output_path = os.path.join(self.output_dir, 'bewerbung.docx')
            doc.save(output_path)
            
            logger.info(f"Documents exported to Word: {output_path}")
            return {
                'path': output_path,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            return {
                'path': '',
                'success': False,
                'error': str(e)
            }
    
    def export_to_pdf(self) -> Dict[str, Any]:
        """Export the documents to PDF format."""
        try:
            # First export to Word
            word_result = self.export_to_word()
            if not word_result['success']:
                raise Exception("Word export failed")
            
            # Convert Word to PDF
            word_path = word_result['path']
            pdf_path = os.path.join(self.output_dir, 'bewerbung.pdf')
            
            convert(word_path, pdf_path)
            
            logger.info(f"Documents exported to PDF: {pdf_path}")
            return {
                'path': pdf_path,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            return {
                'path': '',
                'success': False,
                'error': str(e)
            }

def main():
    """Main function to run the document exporter."""
    parser = argparse.ArgumentParser(description='Export documents to various formats')
    parser.add_argument('--format', choices=['word', 'pdf'], required=True,
                      help='Output format (word or pdf)')
    args = parser.parse_args()
    
    exporter = DocumentExporter()
    
    if args.format == 'word':
        result = exporter.export_to_word()
    else:
        result = exporter.export_to_pdf()
    
    if result['success']:
        print(f"Successfully exported documents to {result['path']}")
    else:
        print(f"Error exporting documents: {result.get('error', 'Unknown error')}")
        sys.exit(1)

if __name__ == '__main__':
    main() 
